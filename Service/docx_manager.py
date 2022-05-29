import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def to_docx(md_file):

    docx_file = md_file.replace(".md", ".docx")
    cmdlet = f"pandoc \"{md_file}\" -f markdown -o \"{docx_file}\""
    os.system(cmdlet)

    text = """
    Rouppeplein 16
    1000 Brussel
    Tel. 02 411 29 07
    ondernemingsnr. 0475319893
    RPR BRUSSEL-NEDERLANDSTALIG
    wouter.vandenberge@intecbrussel.be                                
    """
    image = os.path.join(os.getcwd(), "IntecHeader.png")
    # add_header(docx_file, image, text)

    image = os.path.join(os.getcwd(), "IntecFooter.png")

    # add_footer(docx_file, image)


def add_header(docx_file, headerImage, headerText=None):

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        header = document.sections[0].header

        htable = header.add_table(1, 2, Inches(4))
        htab_cells = htable.rows[0].cells
        ht0 = htab_cells[0].add_paragraph()
        kh = ht0.add_run()
        kh.add_picture(headerImage, width=Inches(3))

        if headerText is not None:
            ht1 = htab_cells[1].add_paragraph(headerText)
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # saving the blank document
        document.save(docx_file)


def add_footer(docx_file, footerImage, footerText=None):

    document = None

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        footer = document.sections[0].footer

        ftable = footer.add_table(1, 2, Inches(4))
        ftab_cells = ftable.rows[0].cells
        ft0 = ftab_cells[0].add_paragraph()
        fh = ft0.add_run()

        imgSize = Inches(4)

        if footerText is not None:
            imgSize = Inches(3)
            ft1 = ftab_cells[1].add_paragraph(footerText)
            ft1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fh.add_picture(footerImage, width=imgSize)

        document.save(docx_file)


if __name__ == "__main__":

    docx_file = "intec_template.docx"

    text = """
    Rouppeplein 16
    1000 Brussel
    Tel. 02 411 29 07
    ondernemingsnr. 0475319893
    RPR BRUSSEL-NEDERLANDSTALIG
    wouter.vandenberge@intecbrussel.be                                
    """

    image = os.getcwd() + os.path.sep + "IntecHeader.png"

    add_header(docx_file, image, text)

    text = "Some footer text.."

    image = os.getcwd() + os.path.sep + "IntecFooter.png"

    add_footer(docx_file, image)
